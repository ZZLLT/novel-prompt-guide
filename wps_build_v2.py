# -*- coding: utf-8 -*-
"""
WPS 网文写作工作台 v2.0
使用 python-docx 直接生成，速度极快
生成后用 WPS 打开
"""
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os
from datetime import datetime

doc = Document()

# ═══════════════════════════════════════════
# 配色与字体
# ═══════════════════════════════════════════
C_PRIMARY = RGBColor(0x1A, 0x36, 0x5D)      # 深蓝
C_PRIMARY_LIGHT = RGBColor(0xEB, 0xF4, 0xFF) # 浅蓝
C_ACCENT = RGBColor(0xC2, 0xA6, 0x00)       # 金橙
C_ACCENT_LIGHT = RGBColor(0xF5, 0xF0, 0xE6) # 浅金
C_TEXT = RGBColor(0x33, 0x33, 0x33)          # 深灰
C_TEXT_LIGHT = RGBColor(0x88, 0x88, 0x88)    # 浅灰
C_BORDER = RGBColor(0xCC, 0xCC, 0xCC)        # 边框灰
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
C_RED = RGBColor(0xCC, 0x00, 0x00)

FONT_TITLE = "黑体"
FONT_HEADING = "微软雅黑"
FONT_BODY = "等线"
FONT_CODE = "Consolas"

# ═══════════════════════════════════════════
# 页面设置
# ═══════════════════════════════════════════
section = doc.sections[0]
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# 默认样式
style = doc.styles['Normal']
style.font.name = FONT_BODY
style.font.size = Pt(11)
style.font.color.rgb = C_TEXT
style.paragraph_format.space_after = Pt(6)
# 设置中文字体
style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)

# ═══════════════════════════════════════════
# 构建函数
# ═══════════════════════════════════════════

def set_cn_font(run, font_name):
    """设置中文字体"""
    run.font.name = font_name
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}></w:rFonts>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)

def add_para(text, font_name=FONT_BODY, size=11, bold=False, color=C_TEXT,
             alignment=None, space_after=6, space_before=0, italic=False):
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if text:
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.color.rgb = color
        run.font.italic = italic
        set_cn_font(run, font_name)
    return p

def add_heading(text, level=1):
    p = doc.add_paragraph()
    sizes = {1: 18, 2: 14, 3: 12}
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.size = Pt(sizes.get(level, 14))
    run.font.bold = True
    run.font.color.rgb = C_PRIMARY
    set_cn_font(run, FONT_HEADING)
    return p

def add_divider(char="─", count=55, color=C_BORDER, size=8):
    add_para(char * count, FONT_BODY, size, color=color, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4, space_before=4)

def add_accent_divider():
    add_para("━" * 55, FONT_BODY, 10, color=C_ACCENT, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6, space_before=6)

def add_blank(n=1):
    for _ in range(n):
        add_para("", space_after=4)

def add_page_break():
    doc.add_page_break()

def add_table(headers, rows, col_widths=None):
    """创建格式化表格"""
    ncols = len(headers)
    nrows = len(rows) + 1
    table = doc.add_table(rows=nrows, cols=ncols, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(h))
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = C_PRIMARY
        set_cn_font(run, FONT_HEADING)
        # 底色
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="EBF4FF" w:val="clear"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    # 数据行
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            if c < ncols:
                cell = table.rows[r + 1].cells[c]
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(str(val))
                run.font.size = Pt(10)
                run.font.color.rgb = C_TEXT
                set_cn_font(run, FONT_BODY)

    # 列宽
    if col_widths:
        for i, w in enumerate(col_widths):
            if i < ncols:
                for row in table.rows:
                    row.cells[i].width = Cm(w)

    add_blank()
    return table

def add_prompt_box(title, content):
    add_para(f"📋 {title}", FONT_HEADING, 11, True, C_ACCENT, space_after=3, space_before=8)
    p = add_para(content, FONT_BODY, 9, color=C_TEXT_LIGHT, italic=True, space_after=6)
    # 边框和底色
    pPr = p._p.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F0E6" w:val="clear"/>')
    pPr.append(shd)
    for border_name in ['top', 'left', 'bottom', 'right']:
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:{border_name} w:val="single" w:sz="4" w:space="4" w:color="C2A600"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)

# ═══════════════════════════════════════════
# 页眉页脚
# ═══════════════════════════════════════════
header = doc.sections[0].header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
hr = hp.add_run("网文AI写作工作台")
hr.font.size = Pt(9)
hr.font.color.rgb = C_TEXT_LIGHT
set_cn_font(hr, FONT_BODY)

footer = doc.sections[0].footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run("《作品名称》 | 第 ")
fr.font.size = Pt(9)
fr.font.color.rgb = C_TEXT_LIGHT
set_cn_font(fr, FONT_BODY)
# 页码
fr2 = fp.add_run("1")
fr2.font.size = Pt(9)
fr2.font.color.rgb = C_TEXT_LIGHT
fr3 = fp.add_run(" 页")
fr3.font.size = Pt(9)
fr3.font.color.rgb = C_TEXT_LIGHT
set_cn_font(fr3, FONT_BODY)

# ═══════════════════════════════════════════
# 封面
# ═══════════════════════════════════════════
for _ in range(8):
    add_blank()

add_accent_divider()
add_para("《　作　品　名　称　》", FONT_TITLE, 28, True, C_PRIMARY, WD_ALIGN_PARAGRAPH.CENTER, 8)
add_para("Novel Writing Workstation v2.0", FONT_HEADING, 13, color=C_TEXT_LIGHT, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
add_accent_divider()

add_blank(2)

# 封面信息表
info_headers = ["项目", "填写", "参考"]
info_rows = [
    ["📖 类型", "______________", "玄幻 / 都市 / 科幻 / 历史 / 悬疑 / 言情..."],
    ["✍️ 风格", "______________", "爽文快节奏 / 慢热铺陈 / 张弛有度 / 偏文艺 / 硬核设定"],
    ["📅 日期", datetime.now().strftime("%Y年%m月%d日"), ""],
    ["💡 核心创意", "______________", "一句话说出你的故事最吸引人的点"],
    ["🎯 目标字数", "______________", "例如：100万字 / 300万字"],
]
add_table(info_headers, info_rows, [2.8, 5, 8])

add_blank(4)
add_para("「 用AI写出你的故事 」", FONT_HEADING, 12, color=C_TEXT_LIGHT, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)

add_page_break()

# ═══════════════════════════════════════════
# 使用指南
# ═══════════════════════════════════════════
add_heading("📖 使用指南", 1)
add_para("本文档是你的网文写作工作台。每个模块包含三层结构：", FONT_HEADING, 11, True)
add_blank()

add_table(
    ["层级", "内容", "操作"],
    [
        ["💭 引导问题", "帮你思考该写什么", "阅读后自行思考或跳过"],
        ["✏️ 编辑区", "表格 / 空白段落", "直接在里面打字填写"],
        ["📋 AI提示词", "可直接使用的AI指令", "选中 → 复制 → 发送给AI → 粘贴回来"],
    ],
    [2.5, 5, 7.5]
)

add_para("工作流程：世界观 → 人物 → 剧情 → 章节 → 润色", FONT_HEADING, 11, True, C_PRIMARY)
add_para("每完成一个阶段打钩 ☐，随时看到进度。", FONT_BODY, 11, color=C_TEXT_LIGHT)
add_accent_divider()

# 进度
add_heading("📊 项目进度总览", 2)
add_table(
    ["🌍 世界观", "👤 人物", "📋 剧情", "✍️ 章节", "🔧 润色", "✅ 完成"],
    [["☐", "☐", "☐", "☐", "☐", "☐"]],
    [2.5, 2.5, 2.5, 2.5, 2.5, 2.5]
)

add_page_break()

# ═══════════════════════════════════════════
# 阶段1：世界观
# ═══════════════════════════════════════════
add_heading("🌍 第一阶段：世界观构建", 1)
add_para("目标：定义世界的基本规则，让后续所有剧情有据可依。", FONT_HEADING, 11, True)
add_para("「世界不是背景板，是所有冲突的土壤」", FONT_BODY, 11, color=C_TEXT_LIGHT)
add_accent_divider()

# 1.1
add_heading("1.1 世界基本设定", 2)
add_table(
    ["设定维度", "你的设定"],
    [["世界类型", ""], ["时代背景", ""], ["世界规模", ""],
     ["核心冲突", ""], ["科技水平", ""],
     ["独特规则(1)", ""], ["独特规则(2)", ""]],
    [3, 12]
)
add_prompt_box("世界观基本设定", """你是一位专业的网络小说世界观架构师。请为一部[类型]小说构建世界的基本框架：

【请回答以下问题，每项100-200字】
1. 这是一个什么样的世界？它的独特之处是什么？
2. 时代背景和社会氛围如何？
3. 世界面临的最大矛盾是什么？（这决定了故事的核心冲突）
4. 普通人不知道但确实存在的「隐藏真相」是什么？
5. 说出这个世界最吸引读者的3个独特设定。""")

# 1.2 力量体系
add_heading("1.2 力量/能力体系", 2)
add_table(
    ["等级", "名称", "标志性能力", "代表人物示例", "备注"],
    [["Lv.1","","","",""],["Lv.2","","","",""],["Lv.3","","","",""],
     ["Lv.4","","","",""],["Lv.5","","","",""],["Lv.6","","","",""],
     ["Lv.7","","","",""],["Lv.8","","","",""],["Lv.9","","","",""],
     ["Lv.10","","","",""]],
    [1.2, 2, 4, 3, 4.8]
)
add_table(
    ["力量维度", "设定"],
    [["力量来源", ""], ["提升方式", ""], ["限制/代价", ""], ["克制关系", ""]],
    [3, 12]
)
add_prompt_box("力量体系设计", """请设计一个完整的[类型]世界观力量体系：

1. 力量来源的创新机制（不要套用老套路，给出新解释）
2. 完整的10级等级体系，每级包含：名称 / 核心能力 / 标志特征
3. 升级条件与瓶颈设计
4. 力量的代价和限制（越强越危险？有使用次数？）
5. 属性克制关系图
6. 主角在这个体系中的「特殊优势」是什么？

格式要求：用清晰的等级表格呈现。""")

# 1.3 势力
add_heading("1.3 势力与组织", 2)
add_table(
    ["势力名称", "类型", "宗旨/目标", "实力定位", "与主角关系"],
    [["", "官方/民间/地下/跨国", "", "S/A/B/C级", ""]] + [["","","","",""] for _ in range(5)],
    [3, 2.5, 4, 2, 3.5]
)
add_prompt_box("势力组织设计", """请设计小说世界中的5-7个主要势力：

每个势力包含：名称 / 类型 / 领袖 / 核心成员 / 资源 / 与主角的利害关系
其中至少要有：
- 1个主角所属或亲近的组织
- 1个直接对抗的敌对势力
- 1个中立但随时可能倒向任何一方的势力
- 1个隐藏极深的幕后组织

给出「势力关系图」文字描述：谁和谁结盟？谁在利用谁？""")

# 1.4 地理
add_heading("1.4 地理与场景", 2)
add_table(
    ["地点名称", "类型", "核心特点", "故事作用"],
    [["","城市/秘境/基地/遗迹","",""]] + [["","","",""] for _ in range(4)],
    [3, 2.5, 5, 4.5]
)

# 1.5 历史
add_heading("1.5 历史与秘密", 2)
add_table(
    ["事件/秘密", "时间", "对现在的影响"],
    [["","",""]] + [["","",""] for _ in range(3)],
    [6, 3, 6]
)
add_prompt_box("历史与隐藏真相", """请构建世界的历史暗线：
1. 列出3个影响当前格局的「大事件」
2. 至少1个「官方历史是谎言」的真相
3. 1个关于世界的终极秘密（到故事后期才揭晓）
4. 主角的能力/身份和这些秘密有什么关联？""")

add_page_break()

# ═══════════════════════════════════════════
# 阶段2：人物
# ═══════════════════════════════════════════
add_heading("👤 第二阶段：人物设计", 1)
add_para("目标：创造读者会记住的角色。主角要立体，反派要有逻辑，配角要有存在感。", FONT_HEADING, 11, True)
add_accent_divider()

add_heading("2.1 主角档案", 2)
add_table(
    ["维度", "项目", "设定"],
    [["基本信息","姓名及含义",""],["","年龄 / 性别",""],
     ["外貌特征","身高 / 体型",""],["","容貌特点(2-3个标志)",""],["","常服/战斗服风格",""],
     ["身份背景","出身 / 家庭",""],["","公开身份",""],["","隐藏身份",""],
     ["性格特质","核心性格(2词概括)",""],["","3个优点+表现",""],["","2个缺陷+剧情影响",""],
     ["","口头禅 / 习惯动作",""],
     ["价值观","底线和软肋",""]],
    [2, 4, 9]
)
add_table(
    ["能力维度", "项目", "设定"],
    [["金手指","能力名称与来源",""],["","详细机制说明",""],["","限制与代价",""],["","成长路径",""],
     ["战斗风格","定位与打法",""],
     ["成长弧线","开篇→中期→结局",""]],
    [2, 4, 9]
)
add_prompt_box("主角设计", """请设计一位[类型]小说的主角，要求角色立体、有血有肉：

1. 【基本信息与外貌】要有画面感，读者能"看见"这个人
2. 【身份背景】出身 + 3个成长关键事件 + 公开/隐藏身份
3. 【性格系统】核心性格 / 3优点3缺陷 / 行为习惯 / 口头禅
4. 【金手指】详细机制 + 限制 + 代价 + 成长路径（这是最核心的！）
5. 【动机系统】终极目标 / 驱动力量 / 价值观 / 底线
6. 【成长弧线】开篇状态 → 中期转变 → 结局理想形态
7. 【读者吸引点】为什么读者会喜欢/代入这个角色？写出3个理由""")

add_heading("2.2 重要角色卡片", 2)
add_para("【女主角 / 最重要的异性角色】", FONT_HEADING, 11, True)
add_table(
    ["项目", "设定"],
    [["姓名 / 年龄",""],["外貌特征",""],["性格(3词)",""],["特殊能力",""],
     ["与主角关系发展",""],["角色弧线",""],["经典台词构思",""]],
    [3.5, 11.5]
)
add_para("【主要反派】", FONT_HEADING, 11, True)
add_table(
    ["项目", "设定"],
    [["姓名 / 年龄",""],["身份 / 势力",""],
     ["核心理念(他为什么觉得自己是对的？)",""],["能力与实力",""],
     ["与主角冲突根源",""],["角色魅力(反派不是纯恶)",""],["最终结局构想",""]],
    [5, 10]
)
add_para("【配角速写表】（每人一行，快速定型）", FONT_HEADING, 11, True)
add_table(
    ["角色", "定位", "核心特质", "与主角关系", "剧情作用"],
    [["导师/前辈","","","",""],["伙伴A","","","",""],["伙伴B","","","",""],
     ["搞笑担当","","","",""],["亦敌亦友","","","",""],["情报/后勤","","","",""]],
    [2, 2, 4, 3, 4]
)
add_prompt_box("配角群像", """请设计以下角色群像：

1. 【女主角】性格独立、有自己追求，不是工具人。详细设定+与主角的关系发展线
2. 【反派】要有魅力的反派！写出他的动机逻辑，让读者觉得"他说得也有道理"但依然站在主角这边
3. 【导师】教主角什么？结局如何？（最好是有牺牲或反转）
4. 【2个伙伴】各有独特技能+性格缺陷，互为补充
5. 【搞笑角色】调节气氛但不喧宾夺主，有自己的"高光时刻"
6. 【亦敌亦友】和主角竞争又合作，关系复杂有张力
7. 给3-5个具体的"角色互动名场面"构思""")

add_heading("2.3 人物关系图谱", 2)
add_table(
    ["角色A", "角色B", "关系性质", "亲密(1-10)", "信任(1-10)"],
    [["主角","女主","","",""],["主角","反派","","",""],["主角","导师","","",""],
     ["主角","伙伴A","","",""],["主角","伙伴B","","",""],
     ["女主","反派","","",""],["","","","",""]],
    [2.5, 2.5, 3, 2, 2]
)

add_page_break()

# ═══════════════════════════════════════════
# 阶段3：剧情
# ═══════════════════════════════════════════
add_heading("📋 第三阶段：剧情大纲", 1)
add_para("目标：把故事从「模糊的想法」变成「可执行的蓝图」。", FONT_HEADING, 11, True)
add_accent_divider()

add_heading("3.1 主线结构", 2)
add_para("【一句话概括】", FONT_HEADING, 11, True)
add_para("这是一个关于________________的故事，主角________________，最终________________。", color=C_TEXT_LIGHT)
add_blank()
add_table(
    ["结构", "内容"],
    [["第一幕 (开篇)",""],["触发事件",""],["第一转折点",""],
     ["第二幕 (对抗)",""],["中点转折",""],["第三幕 (结局)",""]],
    [3, 12]
)
add_prompt_box("主线剧情", """请根据已有的世界观和人物，设计完整的主线剧情：

1. 一句话故事梗概
2. 三幕式详细展开（每幕300-500字）
3. 5个重要剧情转折点（每个要有"读者没想到但合情合理"的效果）
4. 核心冲突的3个层级：外在冲突 / 内在冲突 / 人际冲突
5. 结局的3种可能方向，选最优方案并说明理由""")

add_heading("3.2 分卷规划", 2)
add_table(
    ["卷", "卷名", "一句话概要", "主角目标", "核心冲突", "关键事件数", "预估字数"],
    [["一","","","","","5-8","15-20万"],["二","","","","","5-8","15-25万"],
     ["三","","","","","5-8","20-30万"],["四","","","","","5-8","20-30万"],
     ["五","","","","","5-8","20-30万"],["六","","","","","5-8","15-25万"]],
    [1, 2, 3, 2.5, 2.5, 1.5, 1.5]
)
add_prompt_box("分卷规划", """请将小说分为4-6卷，每卷详细规划：

每卷包含：
1. 卷名（要吸引人！）
2. 本卷一句话概述
3. 主角本卷的成长目标
4. 5-7个关键事件列表
5. 本卷最高潮的场景描述
6. 卷末钩子（让读者必须看下一卷）
7. 预估字数范围
8. 本卷引入的新角色/新设定""")

add_heading("3.3 爽点地图与钩子策略", 2)
add_para("【爽点类型清单】（勾选你计划使用的）", FONT_HEADING, 11, True)
add_table(
    ["爽点类型", "计划在哪卷/哪章使用"],
    [["☐ 扮猪吃虎",""],["☐ 逆袭打脸",""],["☐ 机缘获得",""],
     ["☐ 身份揭露",""],["☐ 越级战斗",""],["☐ 知识碾压",""],
     ["☐ 势力归顺",""],["☐ 前敌变盟友",""]],
    [3.5, 11.5]
)
add_para("【章末钩子策略】（至少5种）", FONT_HEADING, 11, True)
add_table(
    ["钩子类型", "具体使用方式"],
    [["信息悬念(揭示秘密一角)",""],["危机悬念(危险突然降临)",""],
     ["人物悬念(重要角色突然出现/离开)",""],["反转悬念(读者以为的其实是错的)",""],
     ["实力悬念(主角展示新能力的预告)",""],["自定义：",""]],
    [5, 10]
)

add_heading("3.4 伏笔追踪表", 2)
add_table(
    ["编号", "伏笔内容", "埋设位置", "揭示位置", "揭晓方式"],
    [["F01","","第__章","第__章",""],["F02","","第__章","第__章",""],
     ["F03","","第__章","第__章",""],["F04","","第__章","第__章",""],
     ["F05","","第__章","第__章",""],["F06","","第__章","第__章",""],
     ["F07","","第__章","第__章",""]],
    [1, 5, 2, 2, 5]
)

add_page_break()

# ═══════════════════════════════════════════
# 阶段4：章节
# ═══════════════════════════════════════════
add_heading("✍️ 第四阶段：章节写作", 1)
add_para("目标：按计划产出正文。用AI提示词写每一章，用追踪表管理进度。", FONT_HEADING, 11, True)
add_accent_divider()

add_heading("4.1 章节进度追踪表", 2)
add_table(
    ["章", "标题", "类型", "字数", "状态", "爽点", "备注"],
    [["1","","开局","","☐","",""],["2","","","","☐","",""],
     ["3","","","","☐","",""],["4","","","","☐","",""],
     ["5","","小高潮","","☐","",""],["6","","","","☐","",""],
     ["7","","","","☐","",""],["8","","","","☐","",""]],
    [0.8, 2.5, 1.5, 1, 1, 3, 5]
)
add_para("状态标注：☐未开始 / ◐进行中 / ●已完成 / ★需修改", size=9, color=C_TEXT_LIGHT)

add_heading("4.2 单章写作模板", 2)
add_para("以下每写一章可以复制一份模板。", color=C_TEXT_LIGHT)
add_accent_divider()
add_para("第____章：________________", FONT_HEADING, 14, True, C_PRIMARY, space_after=4)
add_table(
    ["项目", "内容"],
    [["本章目标",""],["核心看点/爽点",""],["出场角色",""],
     ["承接上章",""],["本章伏笔(新埋)",""],["章末钩子",""]],
    [3, 12]
)
add_para("【正文】", FONT_HEADING, 12, True)
for _ in range(12):
    add_blank()
add_table(
    ["自查项", "✓"],
    [["开头第一段够吸引吗？","☐"],["结尾有没有钩子？","☐"],
     ["爽点到位了吗？","☐"],["人物性格一致吗？","☐"]],
    [6, 9]
)
add_prompt_box("章节写作提示词", """你是一位专业的网络小说写手。请根据以下信息撰写第X章：

[将上面填好的章节信息粘贴在此处]

写作要求：
- 字数：2000-4000字
- 第一段就要抓住读者（一个悬念/一个冲突/一个画面）
- 对话和动作描写交替，节奏紧凑
- 心理描写适度，不要大段独白
- 结尾必须有明确的钩子
- 保持人物性格和世界设定的一致性""")

add_heading("4.3 批量章节大纲生成", 2)
add_table(
    ["章", "标题", "一句话概括", "核心看点", "出场角色", "章末钩子"],
    [["__","","","","",""] for _ in range(10)],
    [1, 2.5, 3.5, 3, 2.5, 2.5]
)
add_prompt_box("批量大纲生成", """请为第X章到第Y章生成详细大纲。每章包含：
1. 章节标题（要吸引人）
2. 一句话概括
3. 核心看点/冲突
4. 主要出场角色
5. 章节结尾钩子
6. 本章情感走向（↑/↓/↗/↘/→）

要求：章与章之间逻辑递进，高潮有铺垫，日常有意义。""")

add_page_break()

# ═══════════════════════════════════════════
# 阶段5：润色
# ═══════════════════════════════════════════
add_heading("🔧 第五阶段：润色修改", 1)
add_para("目标：提升已完成内容的品质。用AI帮你做专业编辑的工作。", FONT_HEADING, 11, True)
add_accent_divider()

editing_modules = [
    ("5.1 设定一致性检查", """你是一位专业的网文编辑。请检查以下章节是否存在设定不一致：

检查清单：
☐ 人物能力等级前后一致吗？
☐ 世界观规则有没有被违反？
☐ 时间线合理吗？
☐ 人物行为符不符合性格设定？
☐ 物品/道具的出现和消失有没有交代？

请在检查后列出所有问题，按严重程度排序，给出具体修改建议。

[将章节正文粘贴在此]"""),
    ("5.2 文笔逐段润色", """你是一位专业的文字编辑。请逐段润色以下内容：

润色方向：
☐ 提升画面感（让读者"看到"场景）
☐ 优化句子节奏（长短句交替，避免单调）
☐ 精简冗余（删掉废话和水字数的段落）
☐ 强化情感（让情绪传达更到位）
☐ 优化对话（让每句话都推动剧情或展示人物）

要求：保持作者原意和风格，标记主要修改处。

[将章节正文粘贴在此]"""),
    ("5.3 钩子与爽点审计", """你是一位网文阅读体验分析师。请分析以下章节的"阅读体验"：

分析维度：
☐ 第1段：能在3秒内抓住读者吗？不能的话给出改写建议
☐ 最后1段：钩子有力吗？没有的话写出3个钩子方案
☐ 爽点定位：本章有爽点吗？在哪里？够不够爽？
☐ 水字数量化：哪些段落可以删掉或压缩？
☐ 节奏诊断：哪里太慢？哪里太快？
☐ 弃书风险评估：读者最可能在哪一段弃书？

[将章节正文粘贴在此]"""),
    ("5.4 对话专项诊断", """你是一位对话写作专家。请诊断以下对话部分：

诊断项：
☐ 每句话是否体现说话人性格？（还是所有人都一个调调？）
☐ 对话是否推动了剧情或展示了人物关系？
☐ 有没有"为了对话而对话"的废话？
☐ 潜台词丰富吗？（好的对话是冰山，水面下还有9/10）
☐ 信息密度合适吗？

给出每段对话的具体评分和改进建议。

[将对话内容粘贴在此]"""),
]

for title, prompt in editing_modules:
    add_heading(title, 2)
    add_para("【诊断结果记录区】", FONT_HEADING, 11, True)
    add_blank(3)
    add_divider()
    add_prompt_box(f"AI编辑 - {title}", prompt)
    add_blank()

add_page_break()

# ═══════════════════════════════════════════
# 附录
# ═══════════════════════════════════════════
add_heading("📎 附录", 1)

add_heading("写作日志", 2)
add_table(
    ["日期", "写作内容", "字数", "状态", "遇到的问题", "解决方案", "明日计划"],
    [["","","","","","",""],["","","","","","",""]],
    [2, 3, 1.2, 1.5, 3, 3, 3]
)

add_heading("灵感速记区", 2)
add_para("随时记录闪现的想法：", color=C_TEXT_LIGHT)
for _ in range(6):
    add_blank()

add_heading("写作资源清单", 2)
add_table(
    ["类型", "资源名称", "备注"],
    [["参考书/小说","",""],["参考资料/设定集","",""],
     ["写作工具","",""],["灵感来源","",""],["读者反馈渠道","",""]],
    [3, 6, 6]
)

# ═══════════════════════════════════════════
# 保存并打开
# ═══════════════════════════════════════════
desktop = os.path.expanduser("~/Desktop")
filename = f"网文写作工作台_v2_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
filepath = os.path.join(desktop, filename)
doc.save(filepath)

# Write path
with open(os.path.join(os.path.dirname(__file__), "last_doc_path.txt"), "w", encoding="utf-8") as f:
    f.write(filepath)

print(f"DOC_SAVED:{filepath}")
